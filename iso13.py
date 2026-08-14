import sys
import os
import re
import json
import sqlite3
import subprocess 
from PyQt5.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout,
                             QPushButton, QLabel, QListWidget, QListWidgetItem,
                             QMessageBox, QFileDialog, QTextEdit, QLineEdit,
                             QTableWidget, QToolButton, QStyle, QTableWidgetItem, QComboBox, QDialog,
                             QInputDialog, QScrollArea, QSizePolicy
                            )
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QSize
from PyQt5.QtGui import QColor, QFont, QTextBlockFormat, QIcon, QPixmap

# برای نصب این کتابخانه‌ها: pip install PyQt5 pdfplumber python-docx pandas openpyxl pytesseract Pillow
import pdfplumber
from docx import Document
from docx.shared import Inches 
import pandas as pd
import pytesseract
from PIL import Image

HELP_TEXT = """
# راهنمای استفاده از برنامه ممیزی ISO 45001
این راهنما برای کمک به شما در نصب، راه‌اندازی و استفاده از برنامه ممیزی سیستم مدیریت ایمنی و بهداشت شغلی ISO 45001 طراحی شده است. لطفاً تمام مراحل را به دقت دنبال کنید.
1.	پیش‌نیازهای سیستم
قبل از اجرای برنامه، مطمئن شوید که موارد زیر روی سیستم شما نصب شده‌اند:
این برنامه با پایتون نسخه 3 توسعه داده شده است. توصیه می‌شود از آخرین نسخه‌های پایدار پایتون 3 استفاده کنید. می‌توانید آن را از وب‌سایت رسمی پایتون دانلود و نصب کنید [python.org](https://www.python.org/)
*نکته مهم در هنگام نصب پایتون: حتماً گزینه "Add Python to PATH" را در مراحل اولیه نصب فعال کنید تا دسترسی به پایتون از طریق خط فرمان (Command Prompt/Terminal) آسان شود.
Tesseract OCR  اختیاری اما توصیه شده برای تصاویر:
این برنامه قابلیت استخراج متن از فایل‌های تصویری (PNG, JPG و...) را از طریق تکنولوژی OCR (بازشناسی نوری نویسه‌ها) دارد. برای فعال شدن این قابلیت، باید Tesseract OCR را روی سیستم خود نصب کنید.
*نصب بسته‌های زبان فارسی و انگلیسی: در طول نصب  Tesseract، حتماً بسته‌های زبان فارسی (Persian) و انگلیسی (English) را انتخاب و نصب کنید.
اگر Tesseract را نصب نکرده‌اید،. در این صورت، برنامه قادر به استخراج متن از تصاویر نخواهد بود و در صورت تلاش برای پردازش فایل تصویری، پیغام خطا نمایش می‌دهد.

2.	 نصب کتابخانه‌های پایتون
این برنامه از چندین کتابخانه پایتون برای کار با رابط کاربری (GUI)، پایگاه داده و پردازش اسناد استفاده می‌کند. تمام این کتابخانه‌ها را می‌توانید با استفاده از `pip` (مدیر بسته پایتون) نصب کنید.
*خط فرمان (Command Prompt یا Terminal) را باز کنید.
*دستور زیر را اجرا کنید تا تمامی کتابخانه‌های مورد نیاز نصب شوند:
    pip install PyQt5 pdfplumber python-docx pandas openpyxl pytesseract Pillow
*توضیح کتابخانه‌ها:
`PyQt5`: برای ساخت رابط کاربری گرافیکی (GUI) برنامه.
`pdfplumber`: برای استخراج متن از فایل‌های PDF.
 `python-docx`: برای خواندن و پردازش فایل‌های Microsoft Word (DOCX).
`pandas`: برای خواندن و پردازش فایل‌های Excel (XLS, XLSX).
`openpyxl`: یک پیش‌نیاز برای `pandas` برای کار با فرمت‌های Excel.
`pytesseract`: واسط پایتون برای موتور Tesseract OCR (برای استخراج متن از تصاویر).
`Pillow` (PIL): کتابخانه پردازش تصویر پایتون، که توسط `pytesseract` برای باز کردن و آماده‌سازی تصاویر استفاده می‌شود.

3.	نکات مهم در استفاده از برنامه
*پایگاه داده (audit_data.db): برنامه به صورت خودکار یک فایل پایگاه داده SQLite با نام `audit_data.db` را در کنار فایل اجرایی برنامه ایجاد می‌کند. این پایگاه داده حاوی اطلاعات بندهای استاندارد ISO 45001 و شواهد مورد انتظار است. نیازی به ایجاد یا ویرایش دستی این فایل نیست.
*استخراج محتوا از فایل‌ها:
 برنامه می‌تواند متن را از فرمت‌های `PDF`, `DOCX`, `XLS`, `XLSX`, `TXT` و `تصاویر (PNG, JPG, JPEG, TIFF, BMP)` استخراج کند.
 برای فایل‌های تصویری، همانطور که در بخش پیش‌نیازها گفته شد، `Tesseract OCR` باید نصب و پیکربندی شده باشد.
 در صورت بروز خطا در استخراج متن از فایل‌ها (به خصوص تصاویر)، پیغام‌های خطای مربوطه در برنامه نمایش داده خواهند شد.
"""


class WorkerThread(QThread):
    content_extracted = pyqtSignal(str, str) # file_path, content
    processing_status = pyqtSignal(str, str) # status_message, file_path
    error_occurred = pyqtSignal(str)

    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path

    def run(self):
        self.processing_status.emit("در حال پردازش...", self.file_path)
        try:
            file_extension = os.path.splitext(self.file_path)[1].lower()
            content = None

            if file_extension == ".pdf":
                with pdfplumber.open(self.file_path) as pdf:
                    content = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
            elif file_extension == ".docx":
                doc = Document(self.file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
            elif file_extension in (".xls", ".xlsx"):
                df = pd.read_excel(self.file_path)
                content = df.to_string() 
            elif file_extension in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
                try:
                    
                    if 'pytesseract' in sys.modules and hasattr(pytesseract, 'image_to_string'):
                        img = Image.open(self.file_path)
                        content = pytesseract.image_to_string(img, lang='fas+eng') 
                        if not content.strip(): 
                            self.error_occurred.emit(f"تصویر '{os.path.basename(self.file_path)}' حاوی متن قابل استخراج نیست یا OCR به درستی پیکربندی نشده است.")
                            return
                    else:
                        self.error_occurred.emit(f"OCR برای فایل تصویری '{os.path.basename(self.file_path)}' در دسترس نیست. pytesseract نصب یا پیکربندی نشده است.")
                        return
                except Exception as e:
                    self.error_occurred.emit(f"خطا در پردازش OCR برای فایل '{os.path.basename(self.file_path)}': {e}")
                    return
            elif file_extension == ".txt":
                with open(self.file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            else:
                self.error_occurred.emit(f"نوع فایل '{file_extension}' برای پردازش پشتیبانی نمی‌شود: {os.path.basename(self.file_path)}")
                return

            if content:
                self.content_extracted.emit(self.file_path, content)
                self.processing_status.emit("آماده", self.file_path)
            else:
                self.processing_status.emit("خالی/ناموفق", self.file_path)
                self.error_occurred.emit(f"محتوایی از فایل '{os.path.basename(self.file_path)}' استخراج نشد یا فایل خالی بود.")

        except Exception as e:
            self.error_occurred.emit(f"خطا در پردازش فایل '{os.path.basename(self.file_path)}': {e}")
            self.processing_status.emit("خطا", self.file_path)

class AuditEngine:
    def __init__(self):
       
        self.db_path = "audit_data.db"
        self.conn = None
        self.cursor = None
        self.connect_db() 
        self.setup_database()

    def connect_db(self):
        try:
            self.conn = sqlite3.connect(self.db_path) # استفاده از self.db_path
            self.cursor = self.conn.cursor()
        except sqlite3.Error as e:
            QMessageBox.critical(None, "خطای دیتابیس", f"خطا در اتصال به دیتابیس: {str(e)}")
            sys.exit(1)

    def setup_database(self):
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS iso_standards (
                clause_id INTEGER PRIMARY KEY AUTOINCREMENT,
                clause_number TEXT NOT NULL UNIQUE,
                description TEXT,
                keywords TEXT,
                expected_evidence TEXT
            )
        ''')
        self.conn.commit()

       
        self.cursor.execute("SELECT COUNT(*) FROM iso_standards")
        if self.cursor.fetchone()[0] == 0:
            initial_data = [
                ("4.1: درک سازمان و بستر آن",
                 ["محیط", "بستر", "درونی", "بیرونی", "مسائل", "عوامل", "شرایط", "اهداف استراتژیک", "جهت‌گیری", "سازمان", "OH&S", "سیستم مدیریت", "تحلیل", "شناسایی", "نیروهای", "ارتباط", "هدف", "پیامد", "قابلیت", "تأثیرگذار", "کسب و کار", "اثرات"],
                 "گزارشات تحلیل بستر سازمان (SWOT، PESTEL)، صورت‌جلسات بازنگری مدیریت، تحلیل ذینفعان و نیازها، ثبت مسائل داخلی و خارجی مرتبط با OH&S، نمودار سازمانی، گزارشات ریسک و فرصت استراتژیک، گزارشات هیئت مدیره"),

                ("4.2: درک نیازها و انتظارات طرف‌های ذینفع",
                 ["ذینفعان", "نیاز", "انتظار", "کارکنان", "مراجع قانونی", "مشتریان", "پیمانکاران", "تامین کنندگان", "جامعه", "تعیین", "ارتباط", "محدوده", "مسائل", "ارزیابی", "مستند", "الزامات", "نظرسنجی", "مذاکره", "تعهدات", "نظرات", "مطبوعات"],
                 "لیست ذینفعان و نیازهای آن‌ها، مکاتبات با مراجع قانونی، قراردادها، نظرسنجی از کارکنان، ثبت الزامات مربوط به OH&S از طرف ذینفعان، صورت‌جلسات مشورتی، ارتباطات رسمی با گروه‌های ذینفع، سوابق تحلیل رسانه"),

                ("4.3: تعیین دامنه سیستم مدیریت OH&S",
                 ["دامنه", "محدوده", "مرزها", "کاربرد", "تصمیم گیری", "شامل", "مستثنی", "فعالیت‌ها", "محصولات", "خدمات", "موقعیت‌های فیزیکی", "اختیار", "توانایی کنترل", "برون سپاری", "تعیین", "اعمال", "فرایند", "سایت", "مستند"],
                 "سند دامنه سیستم مدیریت OH&S، نمودار سازمانی، نقشه‌های سایت، توصیف فرآیندها و تعاملات مرتبط با OH&S، لیست محصولات و خدمات، طرح استقرار سیستم مدیریت، توافقنامه دامنه"),

                ("4.4: سیستم مدیریت OH&S",
                 ["سیستم مدیریت", "فرآیندها", "تعاملات", "مستندسازی", "حفظ", "نگهداری", "بهبود مستمر", "اجرا", "توالی", "تعیین", "یکپارچگی", "ساختار", "الزامات استاندارد", "روش", "دستورالعمل", "پیاده سازی", "کارآمد"],
                 "مستندات سیستم مدیریت (کتابچه، روش‌های اجرایی)، نقشه فرآیندها، نمودارهای جریان کار، سوابق ممیزی داخلی سیستم مدیریت، طرح کلی سیستم مدیریت OH&S، برنامه نگهداری سیستم"),

                ("5.1: رهبری و تعهد",
                 ["رهبری", "تعهد", "مدیریت ارشد", "خط مشی", "نقش ها", "مسئولیت", "پاسخگویی", "بهبود مستمر", "حمایت", "اطمینان", "توسعه", "اجرا", "ارتباط", "یکپارچگی", "منابع", "فرهنگ OH&S", "مشارکت", "ابلاغ", "پاسخگویی", "تخصیص", "فعالانه"],
                 "خط‌مشی OH&S امضا شده، صورت‌جلسات بازنگری مدیریت (تأیید تعهد)، تخصیص بودجه و منابع، گزارشات مربوط به فرهنگ ایمنی، نقش‌ها و مسئولیت‌های تعریف شده، برنامه‌های تشویقی ایمنی، بیانیه‌های مدیریت"),

                ("5.2: خط‌مشی OH&S",
                 ["خط مشی", "مناسب", " چارچوب", "تعهد", "مستند", "قابل دسترس", "ارتباط", "بهبود", "حفاظت", "مشورت", "الزامات قانونی", "ریسک", "فرصت", "اجرای", "دسترسی", "قابل درک", "بازنگری", "صلاحیت", "شایستگی", "سلامت شغلی"],
                 "سند خط‌مشی OH&S (امضا شده)، پوسترها، اینترانت، سوابق جلسات توجیهی کارکنان در مورد خط‌مشی، برنامه ارتباطات خط‌مشی، امضای خط‌مشی توسط مدیریت ارشد، سوابق توزیع خط‌مشی"),

                ("5.3: نقش‌های سازمانی، مسئولیت‌ها و اختیارات",
                 ["نقش", "مسئولیت", "اختیار", "سازمانی", "تعیین", "ارتباط", "گزارش دهی", "روابط", "تعیین", "واضح", "قابل درک", "ابلاغ", "نماینده مدیریت", "ارزیابی", "ساختار", "شرح وظایف", "ماتریس"],
                 "چارت سازمانی، شرح وظایف (Job Descriptions)، ماتریس مسئولیت‌ها و اختیارات، فرم‌های تعیین نماینده مدیریت OH&S، ابلاغیه‌ها، ساختار گزارش دهی، ثبت مسئولیت‌های ایمنی"),

                ("5.4: مشارکت و مشورت کارکنان",
                 ["مشارکت", "مشورت", "کارکنان", "فرصت", "بازخورد", "نمایندگان", "آموزش", "اطلاعات", "ریسک", "کنترل", "تغییرات", "عدم انطباق", "عملکرد", "تصمیم‌گیری", "فرآیندها", "موانع", "دسترسی", "ایمنی", "بهداشت", "کمیته"],
                 "صورت‌جلسات کمیته‌های ایمنی و بهداشت، سوابق مشورت با کارکنان (فرم‌ها، نظرسنجی‌ها)، لیست نمایندگان کارکنان، سوابق آموزش مشارکت کارکنان، صندوق پیشنهادات و شکایات، کانال‌های ارتباطی دوسویه، برنامه تشویق مشارکت"),

                ("6.1.1: اقدامات جهت پرداختن به ریسک‌ها و فرصت‌ها (عمومی)",
                 ["ریسک", "فرصت", "شناسایی", "ارزیابی", "کنترل", "اقدامات پیشگیرانه", "خطرات", "صلاحیت", "طرح‌ریزی", "مدیریت", "اهداف", "نتایج مورد انتظار", "یکپارچگی", "فرآیند", "اثربخشی", "پیاده سازی", "حفظ", "سلسله مراتب"],
                 "روش اجرایی شناسایی خطر و ارزیابی ریسک، سوابق ارزیابی ریسک، برنامه‌های اقدام (Action Plans) برای ریسک‌ها و فرصت‌ها، گزارشات اثربخشی اقدامات، گزارشات بازنگری ریسک و فرصت، ثبت اقدامات پیشگیرانه"),

                ("6.1.2: شناسایی خطر و ارزیابی ریسک‌ها و فرصت‌ها",
                 ["شناسایی خطر", "ریسک", "فرصت", "پیشگیرانه", "واکنشی", "تغییرات", "اقدامات", "ارزیابی", "بررسی", "دوره ای", "فرآیندها", "فعالیت‌ها", "روال کار", "منابع", "پتانسیل", "علل", "پیامد", "کنترل", "مستند", "جدید", "رویداد"],
                 "فرم‌های شناسایی خطر، گزارشات ارزیابی ریسک، روش‌های اجرایی کنترل عملیات، لیست کنترل‌های موجود، تحلیل تغییرات در ریسک‌ها و فرصت‌ها، ثبت خطرات و ریسک‌ها، برنامه بازرسی ریسک‌ها"),

                ("6.1.2.1: شناسایی خطر",
                 ["خطر", "شناسایی", "رویداد", "پتانسیل", "خسارت", "موقعیت", "فاکتور", "روال کار", "غیر روال", "اضطراری", "زیرساخت", "تجهیزات", "طراحی", "نیروی انسانی", "دانش", "تکنولوژی", "طراحی", "تغییر", "شبه حادثه", "حادثه", "سابقه", "جمع آوری"],
                 "لیست خطرات شناسایی شده، بازرسی‌های ایمنی و بهداشت، گزارشات حوادث و شبه‌حوادث گذشته، تحلیل گزارشات حوادث، ورودی از کارکنان در مورد خطرات، نقشه‌های HIRA (Hazard Identification Risk Assessment)، چک لیست شناسایی خطر"),

                ("6.1.2.2: ارزیابی ریسک‌ها و فرصت‌های OH&S",
                 ["ارزیابی ریسک", "فرصت", "کنترل", "پذیرش", "پتانسیل", "درجه ریسک", "میزان احتمال", "شدت پیامد", "کنترل‌های موجود", "اثربخشی کنترل‌ها", "ریسک باقیمانده", "فرآیند", "دوره ای", "معیار", "اولویت بندی", "تصمیم گیری", "ماتریس"],
                 "ماتریس ارزیابی ریسک، گزارشات تحلیل ریسک، برنامه‌های کنترل ریسک، سوابق تصمیم‌گیری در مورد پذیرش یا عدم پذیرش ریسک، سوابق ارزیابی اثربخشی کنترل‌ها، گزارشات بازنگری ریسک"),

                ("6.1.2.3: تعیین الزامات قانونی و سایر الزامات",
                 ["الزامات قانونی", "مقررات", "مجوزها", "استاندارد", "سایر الزامات", "شناسایی", "دسترسی", "به روز", "ارزیابی انطباق", "قوانین", "آیین نامه‌ها", "توافقنامه‌ها", "مرتبط", "ثبت", "پایش", "ملی", "بین المللی", "صنعتی", "به روزرسانی"],
                 "لیست جامع الزامات قانونی و سایر الزامات قابل کاربرد، مجوزهای عملیاتی، سوابق ممیزی انطباق قانونی، سوابق به‌روزرسانی الزامات، دسترسی به پایگاه‌های اطلاعاتی قانونی، گزارشات انطباق حقوقی"),

                ("6.1.3: تعیین الزامات قانونی و سایر الزامات",
                 ["الزامات قانونی", "مقررات", "مجوزها", "استاندارد", "سایر الزامات", "شناسایی", "دسترسی", "به روز", "ارزیابی انطباق", "قوانین", "آیین نامه‌ها", "توافقنامه‌ها", "مرتبط", "ثبت", "پایش", "ملی", "بین المللی", "صنعتی", "به روزرسانی"],
                 "لیست جامع الزامات قانونی و سایر الزامات قابل کاربرد، مجوزهای عملیاتی، سوابق ممیزی انطباق قانونی، سوابق به‌روزرسانی الزامات، دسترسی به پایگاه‌های اطلاعاتی قانونی، گزارشات انطباق حقوقی"),

                ("6.1.4: طرح‌ریزی اقدامات",
                 ["طرح‌ریزی", "اقدامات", "یکپارچگی", "اثربخشی", "کنترل", "زمانبندی", "مسئولیت", "منابع", "اجرا", "ارزیابی", "به روزآوری", "ریسک", "فرصت", "اهداف", "الزامات قانونی", "چگونگی", "چه کسی", "چه زمانی"],
                 "برنامه‌های اقدام برای پرداختن به ریسک‌ها و فرصت‌ها، جدول زمانبندی، تعیین مسئولیت‌ها و منابع مربوط به اقدامات، سوابق پیگیری و اثربخشی اقدامات، گزارشات پیشرفت اقدامات، بودجه بندی اقدامات"),

                ("6.2.1: اهداف OH&S",
                 ["هدف", "ایمنی", "بهداشت", "پایش", "اندازه گیری", "قابل دستیابی", "سازگار", "ارتباط", "چالش برانگیز", "قابل اندازه‌گیری", "مرتبط", "محدود به زمان", "SMART", "مستند", "ابلاغ", "به روزآوری", "عملکرد", "طرح ریزی"],
                 "لیست اهداف OH&S (SMART)، شاخص‌های عملکرد (KPIs) مرتبط، سوابق ابلاغ اهداف به کارکنان، گزارشات وضعیت دستیابی به اهداف، بازنگری اهداف، برنامه دستیابی به اهداف"),

                ("6.2.2: طرح‌ریزی برای دستیابی به اهداف OH&S",
                 ["طرح‌ریزی", "هدف", "اقدامات", "مسئولیت", "منابع", "زمانبندی", "ارزیابی", "چگونگی", "چه کسی", "چه زمانی", "چگونه", "نتایج", "اجرا", "یکپارگگی", "پایش", "اندازه گیری"],
                 "برنامه‌های اجرایی اهداف OH&S، طرح‌های عملیاتی، ماتریس مسئولیت‌ها برای دستیابی به اهداف، سوابق پایش پیشرفت اهداف، بودجه تخصیص یافته برای اهداف، گزارش پیشرفت اهداف"),

                ("7.1: منابع",
                 ["منبع", "نیروی انسانی", "زیرساخت", "محیط عملیات", "تجهیزات", "مالی", "اطلاعات", "تکنولوژی", "دسترسی", "مناسب", "کفایت", "کافی", "تعیین", "فراهم", "تخصیص", "نگهداری", "پشتیبانی"],
                 "تخصیص بودجه برای OH&S، لیست و سوابق نگهداری تجهیزات، سوابق مربوط به زیرساخت‌ها و محیط کار، چارت سازمانی و لیست کارکنان، برنامه تعمیر و نگهداری، ارزیابی کفایت منابع"),

                ("7.2: صلاحیت",
                 ["صلاحیت", "شایستگی", "آموزش", "آگاهی", "مدرک", "مهارت", "ارزیابی", "نیاز", "تعیین", "حفظ", "اثربخشی", "تحصیلات", "تجربه", "ابلاغ", "سوابق", "بازبینی", "مربیگری", "تخصص"],
                 "سوابق آموزشی کارکنان، ماتریس صلاحیت‌ها، گواهینامه‌های آموزشی، سوابق ارزیابی اثربخشی آموزش‌ها و صلاحیت‌های کسب شده، شرح وظایف شامل نیازهای صلاحیت، برنامه توسعه صلاحیت"),

                ("7.3: آگاهی",
                 ["آگاهی", "خط مشی", "مشارکت", "ریسک", "فرصت", "عدم انطباق", "پیامد", "کمک", "مشارکت", "خروج از محیط کار", "نقش", "مسئولیت", "تبادل", "برنامه", "مستند", "ارتباط", "اثر"],
                 "گزارشات جلسات توجیهی در مورد OH&S، محتوای آموزش‌های آگاهی‌بخش (جزوات، اسلایدها)، سوابق مشارکت کارکنان در برنامه‌های ایمنی، آزمون‌های آگاهی سنجی، پوسترها و علائم آگاهی‌بخش"),

                ("7.4: ارتباطات",
                 ["ارتباطات", "درونی", "بیرونی", "چه", "چه کسی", "چه زمانی", "چگونه", "نتیجه", "فرآیند", "به موقع", "پاسخگویی", "مناسب", "قابل فهم", "مستند", "بازخورد", "تبادل", "شفافیت", "اثربخشی"],
                 "روش اجرایی ارتباطات (داخلی و خارجی)، صورت‌جلسات، ایمیل‌های رسمی، تابلو اعلانات، پروتکل‌های ارتباط با مراجع قانونی و ذینفعان، رسانه‌های ارتباطی، سوابق بازخورد ارتباطات"),

                ("7.4.1: عمومی",
                 ["ارتباطات", "نیاز", "چه", "چه زمانی", "چه کسی", "چگونه", "نتیجه", "فرآیند", "ایجاد", "اجرا", "حفظ", "اثربخشی", "موضوع", "هدف", "دوره ای"],
                 "روش اجرایی کلی ارتباطات، تعیین کانال‌های ارتباطی رسمی، سوابق ارتباطات عمومی، ارزیابی اثربربخشی ارتباطات، برنامه ارتباطی"),

                ("7.4.2: ارتباطات داخلی",
                 ["داخلی", "ارتباطات", "کارکنان", "مشارکت", "اطلاعات", "بازخورد", "تبادل", "شفافیت", "صریح", "آسان", "مدیریت", "تغییرات", "نتایج ممیزی", "عدم انطباق", "ایمنی", "بهداشت", "فرآیند"],
                 "صورت‌جلسات داخلی، ایمیل‌های داخلی، پلتفرم‌های ارتباطی داخلی، کانال‌های بازخورد کارکنان، صندوق پیشنهادات، جلسات تیمی، تابلوهای اطلاع‌رسانی، گزارشات داخلی ایمنی"),

                ("7.4.3: ارتباطات خارجی",
                 ["خارجی", "ارتباطات", "ذینفعان", "مراجع قانونی", "پیمانکاران", "مشتریان", "تبادل", "اطلاعات", "شفاف", "رسمی", "شکایات", "بازرسی", "نظارت", "گزارش دهی", "روابط عمومی", "اضطراری"],
                 "مکاتبات با مراجع قانونی، قراردادها با پیمانکاران، گزارشات بازرسی‌های خارجی، سوابق ارتباط با جامعه و همسایگان، پاسخگویی به شکایات ذینفعان خارجی، بیانیه‌های عمومی"),

                ("7.5.1: اطلاعات مستند شده (عمومی)",
                 ["اطلاعات مستند شده", "ایجاد", "به روزآوری", "حفظ", "کنترل", "دسترسی", "مناسبت", "کفایت", "سیستم مدیریت OH&S", "فرمت", "رسانه", "بازنگری", "حفاظت", "امضا", "تایید"],
                 "روش اجرایی کنترل مستندات، لیست مستندات اصلی سیستم مدیریت (Master List)، سوابق بازنگری و توزیع مستندات، نقشه مستندات سیستم، فایل‌های مستند شده"),

                ("7.5.2: ایجاد و به روزآوری",
                 ["ایجاد", "به روزآوری", "مناسبت", "کفایت", "شناسایی", "فرمت", "بازنگری", "تصویب", "کنترل", "ارجاع", "عنوان", "تاریخ", "نویسنده", "شماره ویرایش", "نحوه نگارش", "ابلاغ", "بررسی"],
                 "فرم‌های بازنگری مستندات، سوابق تصویب مستندات توسط افراد مجاز، الگوهای مستندسازی، سوابق نگارش و اصلاح مستندات، لیست مستندات معتبر، سوابق ابلاغ تغییرات مستندات"),

                ("7.5.3: کنترل اطلاعات مستند شده",
                 ["کنترل", "اطلاعات مستند شده", "در دسترس", "مناسب", "محافظت", "توزیع", "دسترسی", "استفاده", "حفاظت", "حفظ", "تصرف", "بایگانی", "محدودیت", "کپی", "بازیابی", "توزیع", "دسترسی", "حفاظت", "امحا", "نسخه"],
                 "لیست توزیع مستندات، سوابق امحای مستندات قدیمی و منسوخ، پروتکل‌های امنیتی اطلاعات، روش اجرایی بایگانی و بازیابی مستندات، کنترل نسخه‌ها، سوابق دسترسی به مستندات"),

                ("8.1.1: طرح‌ریزی و کنترل عملیات (عمومی)",
                 ["عملیات", "طرح‌ریزی", "کنترل", "ریسک", "کنترل‌های عملیاتی", "حذف خطر", "کاهش ریسک", "اجرا", "حفظ", "تداوم", "فرآیندها", "تغییرات", "خرید", "برون سپاری", "طراحی", "مستند", "پیاده سازی", "اندازه گیری", "مدیریت"],
                 "روش‌های اجرایی عملیاتی، دستورالعمل‌های کاری، برنامه‌های تولید/خدمات، سوابق کنترل‌های عملیاتی، گزارشات بازرسی فرآیندها، نقشه فرآیندهای عملیاتی، سوابق عملیاتی روزانه"),

                ("8.1.2: حذف خطرات و کاهش ریسک‌های OH&S",
                 ["حذف خطر", "کاهش ریسک", "کنترل", "سلسله مراتب کنترل", "مهندسی", "مدیریتی", "PPE", "جایگزینی", "اداری", "ایجاد", "پیاده‌سازی", "حفظ", "ریسک باقیمانده", "ارزیابی مجدد", "تغییرات", "برنامه"],
                 "سوابق حذف خطرات، گزارشات کاهش ریسک، لیست PPE (تجهیزات حفاظت فردی) و سوابق توزیع، روش اجرایی مدیریت ریسک، بررسی اثربخشی کنترل‌ها، اقدامات اصلاحی برای کنترل‌ها، تحلیل سلسله مراتب کنترل"),

                ("8.1.3: مدیریت تغییر",
                 ["تغییر", "مدیریت", "کنترل", "دائم", "موقت", "تاثیر", "ریسک", "فرصت", "فناوری", "تجهیزات", "امکانات", "رویه", "نیروی انسانی", "قبل از اجرا", "شناسایی", "برنامه ریزی", "پیاده سازی", "بازنگری", "اعلام", "مجوز"],
                 "روش اجرایی مدیریت تغییر، فرم‌های درخواست تغییر، سوابق تحلیل ریسک و فرصت‌های ناشی از تغییرات، گزارشات اثربخشی تغییرات، کمیته مدیریت تغییر، سوابق اعلام تغییرات به کارکنان"),

                ("8.1.4: برون‌سپاری",
                 ["برون سپاری", "پیمانکار", "کنترل", "مسئولیت", "شایستگی", "سوابق", "پیمان", "فعالیت‌ها", "فرآیندها", "نیاز", "اطمینان", "الزامات", "تضمین", "ارزیابی", "انتخاب", "قرارداد", "نظارت", "تایید"],
                 "روش اجرایی انتخاب و ارزیابی پیمانکاران، قراردادهای پیمانکاری شامل الزامات OH&S، سوابق ممیزی پیمانکاران، ارزیابی عملکرد پیمانکاران، توافقنامه‌های سطح خدمات، گواهینامه‌های پیمانکاران"),

                ("8.1.4.1: تدارکات",
                 ["تدارکات", "خرید", "محصول", "خدمت", "کنترل", "استاندارد", "مشخصات", "سوابق", "الزامات OH&S", "ایمنی", "مواد اولیه", "تجهیزات", "خدمات", "تعیین", "ارتباط", "حفظ", "کالا", "خدمت", "بازرسی", "تحویل", "پذیرش"],
                 "روش اجرایی تدارکات، مشخصات فنی خرید با لحاظ ملاحظات OH&S، سوابق کنترل کیفیت ورودی از منظر ایمنی، لیست تامین‌کنندگان تایید شده، فاکتورهای خرید با جزئیات ایمنی، گزارشات بازرسی اقلام خریداری شده"),

                ("8.1.4.2: پیمانکاران",
                 ["پیمانکار", "برون سپاری", "تضمین", "انطباق", "معیار", "انتخاب", "شایستگی", "کنترل", "هماهنگی", "ارزیابی", "محدودیت", "آسیب", "نظارت", "بازرسی", "صلاحیت", "آموزش", "مجوز", "توافق"],
                 "معیارهای انتخاب و ارزیابی پیمانکاران، سوابق ارزیابی عملکرد پیمانکاران از نظر OH&S، برنامه کنترل پیمانکاران در سایت، گزارشات بازرسی از فعالیت پیمانکاران، سوابق آموزش ایمنی پیمانکاران"),

                ("8.1.4.3: آمادگی در وضعیت اضطراری",
                 ["اضطراری", "آمادگی", "پاسخ", "سناریو", "مانور", "تست", "ارزیابی", "بهبود", "حوادث", "بلايا", "شرایط اضطراری", "آموزش", "برنامه", "تست", "اطلاعات", "ذینفعان", "تجهیزات", "شناسایی", "تعیین", "واکنش"],
                 "طرح واکنش در شرایط اضطراری، سوابق برگزاری مانورهای اضطراری، گزارشات تحلیل حوادث واقعی، سوابق آموزش کارکنان در خصوص واکنش اضطراری، لیست تجهیزات اضطراری، سوابق ارتباط با نهادهای امدادی"),

                ("8.2: آمادگی و واکنش در وضعیت اضطراری",
                 ["اضطراری", "آمادگی", "واکنش", "سناریو", "تمرین", "آموزش", "بازنگری", "شبیه سازی", "حادثه", "بلايا", "برنامه", "پایش", "اقدامات", "رویه", "دوره ای", "ارتباط", "آزمایش", "اطلاعات"],
                 "طرح واکنش در شرایط اضطراری، سوابق مانورهای اضطراری، گزارشات حوادث و شبه‌حوادث، برنامه‌های بازنگری و بهبود طرح‌های اضطراری، سوابق آموزش و آگاهی‌سازی کارکنان در شرایط اضطراری، لیست تجهیزات اضطراری و تاریخ انقضای آنها"),

                ("9.1.1: پایش، اندازه‌گیری، تحلیل و ارزیابی عملکرد (عمومی)",
                 ["پایش", "اندازه گیری", "تحلیل", "ارزیابی", "عملکرد", "شاخص", "اهداف", "آمار", "گزارش", "متن", "موضوع", "مرتبط", "تجهیزات پایش", "نتایج", "الزامات قانونی", "کیفیت داده", "فراوانی", "کنترل", "بازنگری"],
                 "گزارشات پایش و اندازه‌گیری عملکرد OH&S، شاخص‌های عملکرد (KPIs) و نتایج آن‌ها، سوابق کالیبراسیون تجهیزات پایش، تحلیل داده‌های عملکرد، داشبوردهای عملکرد OH&S، تحلیل روند عملکرد"),

                ("9.1.2: ارزیابی انطباق",
                 ["ارزیابی انطباق", "الزامات قانونی", "سایر الزامات", "دوره ای", "بررسی", "ثبت", "قوانین", "مقررات", "مجوزها", "استاندارد", "پایش", "نتایج", "شواهد", "حفظ", "سوابق", "عدم انطباق", "اقدام"],
                 "سوابق ارزیابی انطباق با الزامات قانونی و سایر الزامات، برنامه‌های ممیزی انطباق، گزارشات عدم انطباق قانونی، اقدامات اصلاحی مربوط به انطباق، لیست به‌روز شده الزامات قانونی، سوابق پیگیری انطباق"),

                ("9.2.1: ممیزی داخلی (عمومی)",
                 ["ممیزی داخلی", "برنامه", "دامنه", "تناوب", "روش", "ممیز", "بی طرف", "گزارش", "صلاحیت", "فرآیند", "شواهد", "هدف", "معیار", "یافته", "نتیجه", "تایید", "پیگیری"],
                 "روش اجرایی ممیزی داخلی، برنامه سالانه ممیزی داخلی، گزارشات ممیزی داخلی (شامل یافته‌ها و عدم انطباق‌ها)، لیست ممیزان داخلی صلاحیت‌دار، سوابق آموزش ممیزی داخلی، گواهینامه‌های ممیزان"),

                ("9.2.2: برنامه ممیزی داخلی",
                 ["برنامه ممیزی", "دوره ای", "ریسک", "تغییرات", "اثربخشی", "صلاحیت", "زمانبندی", "معیار", "دامنه", "تناوب", "روش", "فرآیندها", "اهمیت", "مستند", "برنامه ریزی"],
                 "برنامه سالانه ممیزی داخلی، سوابق ریسک‌ها و تغییرات موثر بر برنامه ممیزی، سوابق اصلاح برنامه ممیزی، لیست ممیزان داخلی، سوابق اثربخشی برنامه‌های ممیزی، سوابق بازنگری برنامه ممیزی"),

                ("9.3: بازنگری مدیریت",
                 ["بازنگری مدیریت", "ورودی", "خروجی", "کفایت", "اثربخشی", "تغییرات", "فرصت", "منابع", "اهداف", "عملکرد", "حوادث", "عدم انطباق", "اقدامات اصلاحی", "بهبود مستمر", "نقش ها", "مسئولیت ها", "دوره ای", "تصمیم گیری", "ارزیابی", "برنامه آینده"],
                 "صورت‌جلسات بازنگری مدیریت، گزارشات ورودی و خروجی بازنگری (شامل تصمیمات و اقدامات)، سوابق پیگیری تصمیمات بازنگری مدیریت، نتایج بازنگری مدیریت به کارکنان، برنامه اقدامات پس از بازنگری"),

                ("10.1: عمومی",
                 ["بهبود", "افزایش", "مناسبت", "کفایت", "اثربخشی", "سیستم مدیریت", "عملکرد OH&S", "نتایج مورد انتظار", "پیش فعال", "واکنشی", "فرصت", "شناسایی", "اقدامات", "روند", "برنامه", "پایداری"],
                 "سوابق بهبود مستمر سیستم مدیریت OH&S، برنامه‌های عملیاتی بهبود، گزارشات مربوط به پروژه‌های بهبود، طرح‌های ارتقاء سیستم، نمودارهای روند بهبود"),

                ("10.2: عدم انطباق و اقدام اصلاحی",
                 ["عدم انطباق", "حادثه", "شبه حادثه", "اقدام اصلاحی", "ریشه یابی", "CAPA", "بازنگری", "اصلاح", "علت", "برخورد", "گزارش", "تحقیق", "سوابق", "پیگیری", "اثربخشی", "پتانسیل", "تکرار", "برنامه", "تصحیح", "اقدامات پیشگیرانه"],
                 "سوابق حوادث و شبه‌حوادث، گزارشات عدم انطباق، فرم‌های اقدام اصلاحی (شامل ریشه‌یابی و اثربخشی)، سوابق پیگیری و اثربخشی اقدامات اصلاحی، لیست عدم انطباق‌ها و وضعیت آن‌ها، ثبت ریشه‌یابی حوادث"),

                ("10.3: بهبود مستمر",
                 ["بهبود مستمر", "عملکرد", "سیستم مدیریت", "افزایش", "مناسبت", "اثربخشی", "پیشگیرانه", "پایداری", "بهینه سازی", "فرصت", "ابتکار", "نتایج", "اهداف", "فرآیندها", "فعالیت", "مشارکت", "فرهنگ"],
                 "گزارشات بهبود مستمر، برنامه‌های بهبود سالیانه، سوابق اقدامات اصلاحی و پیشگیرانه کلی، نتایج بازنگری مدیریت و اقدامات ناشی از آن، پیشنهادات بهبود از کارکنان، نمودارهای بهبود عملکرد"),
            ]
            data_to_insert = []
            for item in initial_data:
                full_clause_string = item[0] # "4.1: درک سازمان و بستر آن"
                keywords_list = item[1]     # ["محیط", "بستر", ...]
                expected_evidence_text = item[2] # "گزارشات تحلیل بستر سازمان (...)"

                # Split clause_number from description
                if ":" in full_clause_string:
                    parts = full_clause_string.split(":", 1)
                    clause_number_only = parts[0].strip()
                    description_for_db = parts[1].strip()
                else:
                    clause_number_only = full_clause_string.strip()
                    description_for_db = full_clause_string.strip() # Fallback if no colon

                data_to_insert.append((clause_number_only, description_for_db, json.dumps(keywords_list), expected_evidence_text))

            self.cursor.executemany(
                "INSERT INTO iso_standards (clause_number, description, keywords, expected_evidence) VALUES (?, ?, ?, ?)",
                data_to_insert
            )
            self.conn.commit()

    def load_iso_standards(self):
        """
        Loads all ISO standard clauses from the database.
        Returns a dictionary of clauses.
        """
        self.cursor.execute("SELECT clause_number, description, keywords, expected_evidence FROM iso_standards")
        rows = self.cursor.fetchall()
        standards = {}
        for row in rows:
            clause_number, description, keywords_json, expected_evidence = row
            
            keywords = json.loads(keywords_json)
            standards[clause_number] = {
                "description": description,
                "keywords": keywords,
                "expected_evidence": expected_evidence
            }
        return standards

    def get_all_clauses(self):
        """
        Retrieves all clause numbers, descriptions, keywords, and expected evidence.
        """
        # Corrected table name and primary key column
        self.cursor.execute("SELECT clause_number, description, keywords, expected_evidence FROM iso_standards ORDER BY clause_id")
        clauses = []
        for row in self.cursor.fetchall():
            clause_number, description, keywords_json, expected_evidence = row
            keywords = json.loads(keywords_json) # Keywords are stored as JSON string
            clauses.append((clause_number, description, keywords, expected_evidence))
        return clauses

    def perform_audit(self, documents_content):
        """
        Performs the audit by comparing document content with ISO standards.
        Returns a dictionary of audit results for each clause.
        """
        results = {}
        all_clauses_data = self.load_iso_standards() # Use load_iso_standards to get structured data

       
        combined_content = " ".join(documents_content).lower()
        
        combined_content = re.sub(r'[^ا-یآ-یa-zA-Z\s]', '', combined_content)
        combined_content = re.sub(r'\s+', ' ', combined_content).strip()

        for clause_num, clause_data in all_clauses_data.items():
            clause_description = clause_data["description"]
            required_keywords = clause_data["keywords"]
            expected_evidence_text = clause_data["expected_evidence"]

            clause_non_conformities = []
            clause_found_keywords_overall = set()

            
            for keyword in required_keywords:
                
                if re.search(r'\b' + re.escape(keyword.lower()) + r'\b', combined_content):
                    clause_found_keywords_overall.add(keyword)

            
            if not clause_found_keywords_overall:
                clause_non_conformities.append({
                    "type": "Major (عدم وجود مستندات/پوشش)",
                    "file_path": "مستندات بارگذاری شده کلی", # N/A as it applies to all docs
                    "excerpt": f"هیچ یک از کلمات کلیدی بند '{clause_num} - {clause_description}' در مستندات بارگذاری شده یافت نشد. این بند نیازمند پوشش کامل است.",
                    "found_keywords": [],
                })
            
            elif len(clause_found_keywords_overall) < len(required_keywords) * 0.6:
                clause_non_conformities.append({
                    "type": "Minor (پوشش ناکافی/مبهم)",
                    "file_path": "مستندات بارگذاری شده کلی", # N/A as it applies to all docs
                    "excerpt": f"فقط تعداد محدودی از کلمات کلیدی بند '{clause_num} - {clause_description}' یافت شد ({len(clause_found_keywords_overall)} از {len(required_keywords)}). نیاز به بررسی بیشتر برای اطمینان از پوشش کامل.",
                    "found_keywords": list(clause_found_keywords_overall),
                })
            

            results[clause_num] = {
                "description": clause_description,
                "found_keywords": list(clause_found_keywords_overall),
                "expected_evidence": expected_evidence_text,
                "non_conformities": clause_non_conformities,
                "all_expected_keywords": list(required_keywords) # Adding all expected keywords for better reporting
            }
        return results

    def close_connection(self):
        if self.conn:
            self.conn.close()

class ListItemWidget(QWidget):
    removed = pyqtSignal(str) # سیگنال برای حذف فایل

    def __init__(self, file_path, parent=None):
        super().__init__(parent)
        self.file_path = file_path
        self.init_ui()

    def init_ui(self):
        layout = QHBoxLayout()
        layout.setContentsMargins(5, 5, 5, 5)
        layout.setSpacing(10)

        icon_label = QLabel()
        icon_label.setPixmap(self.get_file_icon(self.file_path).pixmap(QSize(24, 24)))
        layout.addWidget(icon_label)

        self.file_name_label = QLabel(os.path.basename(self.file_path))
        self.file_name_label.setFont(QFont("Tahoma", 9))
        layout.addWidget(self.file_name_label)

        self.status_label = QLabel("منتظر...")
        
        self.status_label.setFont(QFont("Tahoma", 8, QFont.StyleItalic))
        self.status_label.setStyleSheet("color: gray;")
        layout.addWidget(self.status_label)

        layout.addStretch()

        self.remove_button = QToolButton()
        self.remove_button.setIcon(self.style().standardIcon(QStyle.SP_TrashIcon))
        self.remove_button.setToolTip("حذف فایل")
        self.remove_button.clicked.connect(self.remove_self)
        layout.addWidget(self.remove_button)

        self.setLayout(layout)

    def get_file_icon(self, file_path):
        """بر اساس نوع فایل، آیکون مناسب را برمی‌گرداند."""
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == ".pdf":
            return QApplication.instance().style().standardIcon(QStyle.SP_FileIcon) 
        elif file_extension in (".doc", ".docx"):
            return QApplication.instance().style().standardIcon(QStyle.SP_FileIcon) 
        elif file_extension in (".xls", ".xlsx"):
            return QApplication.instance().style().standardIcon(QStyle.SP_FileIcon) 
        elif file_extension in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
            
            return QApplication.instance().style().standardIcon(QStyle.SP_FileIcon) 
            
        elif file_extension == ".txt":
            return QApplication.instance().style().standardIcon(QStyle.SP_FileIcon) 
        else:
            return QApplication.instance().style().standardIcon(QStyle.SP_FileIcon) 
    def set_processing_status(self, status):
        self.status_label.setText(status)
        if status == "در حال پردازش...":
            self.status_label.setStyleSheet("color: orange;")
        elif status == "آماده":
            self.status_label.setStyleSheet("color: green;")
        elif status == "خطا" or status == "خالی/ناموفق":
            self.status_label.setStyleSheet("color: red;")

    def remove_self(self):
        reply = QMessageBox.question(self, "حذف فایل",
                                     f"آیا مطمئنید که می‌خواهید فایل '{os.path.basename(self.file_path)}' را حذف کنید؟",
                                     QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.removed.emit(self.file_path) # سیگنال حذف را ارسال می‌کند

class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.uploaded_files = {}
        self.worker_threads = []
        self.section_file_paths = {}
        self.audit_engine = AuditEngine() 
        self.init_ui()

        welcome_box = QMessageBox(self)
        welcome_box.setWindowTitle(" (Version:01.01.01) نسخه آزمایشی ")
        welcome_box.setText(
            " ISO نرم افزار ممیزی، مدیریت مستندات و راهنمای استقرار 45001:2018 \n\n"   
            "برای راه اندازی سریع برنامه و استفاده بهینه تمامی امکانات برنامه، قبل از ورود فایل راهنما را مطالعه فرمایید\n\n"
            "طراح و برنامه نویس: حسین حجتی سیاح"
        )
        welcome_box.setIcon(QMessageBox.Information)


        ok_button = welcome_box.addButton("ورود به برنامه", QMessageBox.AcceptRole)


        help_button = welcome_box.addButton(" راهنما راه اندازی برنامه", QMessageBox.HelpRole)


        welcome_box.exec_()


        if welcome_box.clickedButton() == help_button:
            self._show_help_dialog()

    def _show_help_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("راهنمای استفاده از برنامه")
        dialog.setMinimumSize(800, 600)

        layout = QVBoxLayout(dialog)

        text_edit = QTextEdit()
        text_edit.setReadOnly(True)
        text_edit.setText(HELP_TEXT)
        text_edit.setFont(QFont("Tahoma", 10))
        text_edit.setTextInteractionFlags(Qt.TextSelectableByMouse | Qt.TextSelectableByKeyboard)
        layout.addWidget(text_edit)

        close_button = QPushButton("بستن")
        close_button.clicked.connect(dialog.accept)
        layout.addWidget(close_button, alignment=Qt.AlignRight)

        dialog.exec_()


        
    def open_file_from_list(self, item):
        file_path = item.data(Qt.UserRole) 
        
        if file_path and os.path.exists(file_path):
            try:
                if sys.platform == "win32":
                    os.startfile(file_path)
                elif sys.platform == "darwin": 
                    subprocess.call(["open", file_path])
                else: # Linux
                    subprocess.call(["xdg-open", file_path])
            except Exception as e:
                QMessageBox.critical(self, "خطا در باز کردن فایل",
                                     f"خطا در باز کردن فایل '{os.path.basename(file_path)}': {e}\n"
                                     "لطفاً اطمینان حاصل کنید که برنامه مناسب برای باز کردن این نوع فایل نصب شده است.")
        else:
            QMessageBox.warning(self, "فایل یافت نشد",
                                 f"مسیر فایل '{os.path.basename(file_path)}' معتبر نیست یا فایل وجود ندارد.")

    def init_ui(self):
        self.setWindowTitle(" ISOسیستم ممیزی 45001:2018")
        self.setMinimumSize(1000, 700)

        self.main_font = QFont("Tahoma", 10)
        self.setFont(self.main_font)

        main_layout = QHBoxLayout()

        
        right_panel = QVBoxLayout()
        right_panel.setContentsMargins(10, 10, 10, 10)
        right_panel.setSpacing(10)
        right_panel.setAlignment(Qt.AlignTop)

        title_files = QLabel("ISOمنو مدیریت مستندات 45001")
        title_files.setFont(QFont("Tahoma", 12, QFont.Bold))
        title_files.setStyleSheet("color: #0056b3;")
        right_panel.addWidget(title_files)

        upload_button = QPushButton("ISOمنو مدیریت مستندات 45001")
        upload_button.setFont(self.main_font)
        upload_button.setStyleSheet("background-color: #4CAF50; color: white; padding: 8px; border-radius: 4px;")
        upload_button.clicked.connect(self.open_sectioned_upload_dialog)
        right_panel.addWidget(upload_button)

        self.file_list_widget = QListWidget()
        self.file_list_widget.setFont(self.main_font)
        self.file_list_widget.setStyleSheet("background-color: #f0f0f0; border: 1px solid #ccc; border-radius: 4px;")
        self.file_list_widget.itemClicked.connect(self.open_file_from_list)
        right_panel.addWidget(self.file_list_widget)

        audit_button = QPushButton(" ISOشروع ممیزی 45001")
        audit_button.setFont(QFont("Tahoma", 11, QFont.Bold))
        audit_button.setStyleSheet("background-color: #007bff; color: white; padding: 10px; border-radius: 4px;")
        audit_button.clicked.connect(self.start_audit)
        right_panel.addWidget(audit_button)

        self.manage_db_button = QPushButton("مدیریت پایگاه دانش نرم افزار")
        self.manage_db_button.setFont(self.main_font)
        self.manage_db_button.setStyleSheet(
            "QPushButton { background-color: #3498db; color: white; padding: 8px; border-radius: 4px; }"
            "QPushButton:hover { background-color: #2980b9; }"
        )
        self.manage_db_button.clicked.connect(self.request_db_password)
        right_panel.addWidget(self.manage_db_button)

        show_all_clauses_button = QPushButton("ISO راهنمای استقرار 45001")
        show_all_clauses_button.setFont(self.main_font)
        show_all_clauses_button.setStyleSheet("background-color: #6c757d; color: white; padding: 8px; border-radius: 4px;")
        show_all_clauses_button.clicked.connect(self.display_all_clauses)
        right_panel.addWidget(show_all_clauses_button)

        main_layout.addLayout(right_panel, 2)

        
        center_panel = QVBoxLayout()
        center_panel.setContentsMargins(10, 10, 10, 10)
        center_panel.setSpacing(10)

        title_results = QLabel("ISO نتایج ممیزی 45001:2018 ")
        title_results.setFont(QFont("Tahoma", 12, QFont.Bold))
        title_results.setStyleSheet("color: #0056b3;")
        center_panel.addWidget(title_results)

        filter_layout = QHBoxLayout()
        filter_label = QLabel("فیلتر نتایج:")
        filter_label.setFont(self.main_font)
        self.filter_combobox = QComboBox()
        self.filter_combobox.setFont(self.main_font)
        self.filter_combobox.addItems(["همه نتایج", "انطباق کامل", "عدم انطباق جزئی", "عدم انطباق اصلی"])
        self.filter_combobox.setStyleSheet("padding: 5px; border: 1px solid #ccc; border-radius: 4px;")
        self.filter_combobox.currentIndexChanged.connect(self.filter_audit_results)
        filter_layout.addWidget(filter_label)
        filter_layout.addWidget(self.filter_combobox)
        filter_layout.addStretch()
        center_panel.addLayout(filter_layout)

        self.results_table = QTableWidget()
        self.results_table.setColumnCount(3)
        self.results_table.setHorizontalHeaderLabels(["بند استاندارد", "وضعیت انطباق", "شواهد مورد انتظار"])
        self.results_table.horizontalHeader().setFont(QFont("Tahoma", 10, QFont.Bold))
        self.results_table.verticalHeader().setVisible(False)
        self.results_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.results_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.results_table.setFont(self.main_font)
        self.results_table.setStyleSheet("background-color: white; border: 1px solid #ccc; border-radius: 4px;")
        self.results_table.horizontalHeader().setStretchLastSection(True)
        self.results_table.setColumnWidth(0, 150)
        self.results_table.setColumnWidth(1, 100)
        self.results_table.setColumnWidth(2, 300)
        self.results_table.cellClicked.connect(self.display_clause_details)
        center_panel.addWidget(self.results_table)

        
        report_button = QPushButton("گزارش گیری")
        report_button.setFont(self.main_font)
        report_button.setStyleSheet("background-color: #28a745; color: white; padding: 8px; border-radius: 4px;")
        report_button.clicked.connect(self.show_report_options)
        center_panel.addWidget(report_button)

        main_layout.addLayout(center_panel, 5)

        
        left_panel = QVBoxLayout()
        left_panel.setContentsMargins(10, 10, 10, 10)
        left_panel.setSpacing(10)
        left_panel.setAlignment(Qt.AlignTop)

        title_details = QLabel("ISO منو راهنمای استقرار هر بند از استاندارد 45001 ")
        title_details.setFont(QFont("Tahoma", 12, QFont.Bold))
        title_details.setStyleSheet("color: #0056b3;")
        left_panel.addWidget(title_details)

        self.clause_details_label = QLabel("جزئیات راهنمایی استقرار بند انتخاب شده.")
        self.clause_details_label.setFont(QFont("Tahoma", 10, QFont.Bold))
        self.clause_details_label.setStyleSheet("color: #0056b3;")
        left_panel.addWidget(self.clause_details_label)

        self.clause_description_text = QTextEdit()
        self.clause_description_text.setReadOnly(True)
        self.clause_description_text.setFont(self.main_font)
        self.clause_description_text.setStyleSheet("background-color: #f8f8f8; border: 1px solid #ddd; border-radius: 4px;")
        self.clause_description_text.setPlaceholderText("توضیحات بند و شواهد مورد انتظار...")
        left_panel.addWidget(self.clause_description_text)

        search_layout = QHBoxLayout()
        self.search_input = QLineEdit()
        self.search_input.setFont(self.main_font)
        self.search_input.setPlaceholderText("جستجو در مستندات بارگذاری شده...")
        self.search_input.setStyleSheet("padding: 5px; border: 1px solid #ccc; border-radius: 4px;")
        self.search_input.returnPressed.connect(self.search_in_content)
        search_button = QPushButton("جستجو")
        search_button.setFont(self.main_font)
        search_button.setStyleSheet("background-color: #007bff; color: white; padding: 5px; border-radius: 4px;")
        search_button.clicked.connect(self.search_in_content)
        search_layout.addWidget(self.search_input)
        search_layout.addWidget(search_button)
        left_panel.addLayout(search_layout)

        self.search_results_text = QTextEdit()
        self.search_results_text.setReadOnly(True)
        self.search_results_text.setFont(self.main_font)
        self.search_results_text.setStyleSheet("background-color: #f8f8f8; border: 1px solid #ddd; border-radius: 4px;")
        self.search_results_text.setPlaceholderText("نتایج جستجو اینجا نمایش داده می‌شود...")
        left_panel.addWidget(self.search_results_text)

        main_layout.addLayout(left_panel, 4)
        
        self.setLayout(main_layout)

        self.current_audit_results = {}
    def open_sectioned_upload_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle(" ISOپنل مدیریت مستندات 45001")
        dialog.setMinimumSize(800, 650)

        main_layout = QVBoxLayout(dialog)
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
    
        section_titles = {
           "4": "بافت سازمان (Context of the organization)",
           "5": "رهبری و مشارکت کارکنان (Leadership and worker participation)",
           "6": "طرح‌ریزی (Planning)",
           "7": "پشتیبانی (Support)",
           "8": "عملیات (Operation)",
           "9": "ارزیابی عملکرد (Performance evaluation)",
           "10": "بهبود (Improvement)"
        }

      

        for clause, title in section_titles.items():
           clause_label = QLabel(f"{clause}. {title}")
           clause_label.setAlignment(Qt.AlignCenter)
           clause_label.setStyleSheet("font-weight: bold; padding: 4px;")
           scroll_layout.addWidget(clause_label)
           
           button_row = QHBoxLayout()
           button_row.setAlignment(Qt.AlignCenter)


           choose_btn = QPushButton("انتخاب فایل(ها)")
           choose_btn.setProperty("clause", clause)
           choose_btn.clicked.connect(lambda _, c=clause: self.select_files_for_clause(c))

           view_btn = QPushButton("مشاهده مدارک")
           view_btn.clicked.connect(lambda _, c=clause: self.view_uploaded_files(c))
           
           button_row.addWidget(view_btn)
           button_row.addWidget(choose_btn)

           scroll_layout.addLayout(button_row)

           spacer = QLabel("")
           spacer.setFixedHeight(8)
           scroll_layout.addWidget(spacer)

        scroll_content.setLayout(scroll_layout)
        scroll_area.setWidget(scroll_content)
        main_layout.addWidget(scroll_area)

        close_btn = QPushButton("بستن")
        close_btn.setFixedWidth(120)
        close_btn.clicked.connect(dialog.accept)

        main_layout.addWidget(close_btn, alignment=Qt.AlignCenter)

        dialog.setLayout(main_layout)
        dialog.exec_()

    def select_files_for_clause(self, clause_number):
        options = QFileDialog.Options()
        file_paths, _ = QFileDialog.getOpenFileNames(self, "انتخاب فایل برای بند " + clause_number, "",
           "All Supported Files (*.pdf *.docx *.xls *.xlsx *.png *.jpg *.jpeg *.tiff *.bmp *.txt)")

        if file_paths:
           if clause_number not in self.section_file_paths:
               self.section_file_paths[clause_number] = []
           existing_paths = [f["path"] for f in self.section_file_paths[clause_number]]
           for path in file_paths:
               self.section_file_paths[clause_number].append({
                   "path": path,
                   "status": "در حال پردازش",
                   "clause": clause_number  
               })
               
               worker = WorkerThread(path)
               worker.content_extracted.connect(self.handle_content_extracted)
               worker.processing_status.connect(self.update_section_file_status)
               worker.error_occurred.connect(self.show_error_message)
               self.worker_threads.append(worker)
               worker.start()
               

           QMessageBox.information(self, "فایل‌ها در حال پردازش‌اند", f"{len(file_paths)} فایل انتخاب شد و در حال پردازش می‌باشند. .")


    def update_section_file_status(self, status, file_path):
        for clause_files in self.section_file_paths.values():
           for file_info in clause_files:
               if file_info["path"] == file_path:
                   file_info["status"] = status

    def view_uploaded_files(self, clause_number):
        files = self.section_file_paths.get(clause_number, [])
        if not files:
           QMessageBox.information(self, "هیچ فایلی وجود ندارد", f"برای بند {clause_number} فایلی انتخاب نشده است.")
           return
        dialog = QDialog(self)
        dialog.setWindowTitle(f"فایل‌های بند {clause_number}")
        dialog.setMinimumSize(600, 400)
        layout = QVBoxLayout(dialog)

        for file_info in files:
           file_path = file_info["path"]
           status = file_info.get("status", "در حال پردازش")

           hbox = QHBoxLayout()


           file_label = QLabel(os.path.basename(file_path))
           file_label.setMinimumWidth(250)
           hbox.addWidget(file_label)


           status_label = QLabel(status)
           color = "red" if status == "بارگذاری شد" else "green"
           status_label.setStyleSheet(f"color: {color}; font-weight: bold;")
           hbox.addWidget(status_label)

           open_btn = QToolButton()
           open_btn.setIcon(self.style().standardIcon(QStyle.SP_DirOpenIcon))
           open_btn.setToolTip("باز کردن فایل")
           open_btn.clicked.connect(lambda _, p=file_path: self.open_file(p))
           hbox.addWidget(open_btn)


           delete_btn = QToolButton()
           delete_btn.setIcon(self.style().standardIcon(QStyle.SP_TrashIcon))
           delete_btn.setToolTip("حذف فایل")
           delete_btn.clicked.connect(lambda _, p=file_path, c=clause_number: self.delete_clause_file(p, c, dialog))
           hbox.addWidget(delete_btn)

           layout.addLayout(hbox)

        close_btn = QPushButton("بستن")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn, alignment=Qt.AlignCenter)

        dialog.setLayout(layout)
        dialog.exec_()
 
    def remove_file(self, file_path):
        for i in range(self.file_list_widget.count()):
            item = self.file_list_widget.item(i)
            widget = self.file_list_widget.itemWidget(item)
            if widget and widget.file_path == file_path:
                self.file_list_widget.takeItem(i)
                del self.uploaded_files[file_path]
                QMessageBox.information(self, "فایل حذف شد", f"فایل '{os.path.basename(file_path)}' با موفقیت حذف شد.")
                return

    def handle_content_extracted(self, file_path, content):
        self.uploaded_files[file_path] = content
        print(f"✅ ذخیره در uploaded_files: {file_path}")
        clause_number = None
        for clause_files in self.section_file_paths.values():
           for file_info in clause_files:
               if file_info["path"] == file_path:
                   file_info["status"] = "بارگذاری شد"
                   clause_number = file_info.get("clause")
                   if not clause_number:
                       clause_number = clause
                   break

        if clause_number:
           print(f"📌 فایل متعلق به بند {clause_number} به‌روزرسانی شد.")
        else:
           print(f"⚠️ بند مربوط به فایل {file_path} پیدا نشد.")

        existing_items = [self.file_list_widget.item(i).data(Qt.UserRole) for i in range(self.file_list_widget.count())]
        if file_path not in existing_items:
           list_item = QListWidgetItem()
           widget = ListItemWidget(file_path)
           widget.removed.connect(self.remove_file)
           widget.set_processing_status("آماده")
           list_item.setSizeHint(widget.sizeHint())
           list_item.setData(Qt.UserRole, file_path)
           self.file_list_widget.addItem(list_item)
           self.file_list_widget.setItemWidget(list_item, widget)

    def show_error_message(self, message):
        QMessageBox.warning(self, "خطا در پردازش", message)
    
    def open_file(self, file_path):
        if os.path.exists(file_path):
           try:
               if sys.platform == "win32":
                   os.startfile(file_path)
               elif sys.platform == "darwin":
                   subprocess.call(["open", file_path])
               else:
                   subprocess.call(["xdg-open", file_path])
           except Exception as e:
               QMessageBox.critical(self, "خطا", f"باز کردن فایل با خطا مواجه شد:\n{e}")
        else:
           QMessageBox.warning(self, "فایل موجود نیست", "مسیر فایل دیگر معتبر نیست.")
 
    def delete_clause_file(self, file_path, clause_number, dialog):
        reply = QMessageBox.question(self, "حذف فایل", f"آیا از حذف '{os.path.basename(file_path)}' مطمئن هستید؟", 
                                     QMessageBox.Yes | QMessageBox.No)
        if reply == QMessageBox.Yes:
           self.section_file_paths[clause_number] = [
               f for f in self.section_file_paths[clause_number] if f["path"] != file_path
           ]
           QMessageBox.information(self, "حذف شد", "فایل با موفقیت حذف شد.")
           dialog.accept()
           self.view_uploaded_files(clause_number)


    def start_audit(self):
        if not self.uploaded_files:
            QMessageBox.warning(self, "مستندات یافت نشد", "لطفاً ابتدا فایل‌های مستندات را بارگذاری کنید.")
            return


        # Wait for all files to be processed
        for file_path, content in self.uploaded_files.items():
            if content is None:
                QMessageBox.warning(self, "پردازش در حال انجام", "لطفاً منتظر بمانید تا تمام فایل‌ها پردازش شوند.")
                return

        documents_content = [content for content in self.uploaded_files.values() if content is not None]
        if not documents_content:
            QMessageBox.warning(self, "محتوایی برای ممیزی وجود ندارد", "هیچ محتوایی از فایل‌های بارگذاری شده استخراج نشد.")
            return

        self.results_table.setRowCount(0)
        self.clause_description_text.clear()
        self.search_results_text.clear()
        self.clause_details_label.setText("جزئیات بند انتخاب شده اینجا نمایش داده می‌شود.")

        self.current_audit_results = self.audit_engine.perform_audit(documents_content)
        self.display_audit_results(self.current_audit_results)
        QMessageBox.information(self, "ممیزی کامل شد", "ممیزی ISO 45001 با موفقیت انجام شد.")

    def display_audit_results(self, results):
        self.results_table.setRowCount(0)

        for clause_number, data in results.items():
            row_position = self.results_table.rowCount()
            self.results_table.insertRow(row_position)

            clause_item = QTableWidgetItem(f"{clause_number} - {data['description']}")
            self.results_table.setItem(row_position, 0, clause_item)

            status_text = "انطباق کامل"
            status_color = QColor("#28a745")
            if data['non_conformities']:
                if any(nc['type'].startswith("Major") for nc in data['non_conformities']):
                    status_text = "عدم انطباق اصلی"
                    status_color = QColor("#dc3545")
                elif any(nc['type'].startswith("Minor") for nc in data['non_conformities']):
                    status_text = "عدم انطباق جزئی"
                    status_color = QColor("#ffc107")

            status_item = QTableWidgetItem(status_text)
            status_item.setForeground(status_color)
            self.results_table.setItem(row_position, 1, status_item)

            
            expected_evidence_text = data.get('expected_evidence', "اطلاعات موجود نیست.")
            expected_evidence_item = QTableWidgetItem(expected_evidence_text)
            self.results_table.setItem(row_position, 2, expected_evidence_item)

        self.results_table.resizeRowsToContents()

    def filter_audit_results(self):
        filter_type = self.filter_combobox.currentText()
        if not self.current_audit_results:
            return

        filtered_results = {}
        for clause_number, data in self.current_audit_results.items():
            status_text = "انطباق کامل"
            if data['non_conformities']:
                if any(nc['type'].startswith("Major") for nc in data['non_conformities']):
                    status_text = "عدم انطباق اصلی"
                elif any(nc['type'].startswith("Minor") for nc in data['non_conformities']):
                    status_text = "عدم انطباق جزئی"

            if filter_type == "همه نتایج":
                filtered_results[clause_number] = data
            elif filter_type == "انطباق کامل" and status_text == "انطباق کامل":
                filtered_results[clause_number] = data
            elif filter_type == "عدم انطباق جزئی" and status_text == "عدم انطباق جزئی":
                filtered_results[clause_number] = data
            elif filter_type == "عدم انطباق اصلی" and status_text == "عدم انطباق اصلی":
                filtered_results[clause_number] = data

        self.display_audit_results(filtered_results)

    def display_clause_details(self, row, column):
        clause_number_full = self.results_table.item(row, 0).text()
        clause_number_only = clause_number_full.split(' ')[0]

       
        conn = sqlite3.connect(self.audit_engine.db_path) 
        cursor = conn.cursor()
        clause_info = cursor.execute(
            "SELECT description, keywords, expected_evidence FROM iso_standards WHERE clause_number = ?",
            (clause_number_only,)
        ).fetchone()
        conn.close()

        if clause_info:
            description, keywords_json, evidence = clause_info
            self.clause_details_label.setText(f"جزئیات بند: {clause_number_full}")

            details_text = f"**توضیحات بند:**\n{description}\n\n"
            details_text += f"**شواهد مورد انتظار:**\n{evidence}"

            self.clause_description_text.setHtml(details_text.replace('\n', '<br>'))
        else:
            self.clause_details_label.setText("جزئیات بند یافت نشد.")
            self.clause_description_text.setHtml("<em>اطلاعات مربوط به این بند در دیتابیس موجود نیست.</em>")

    def display_all_clauses(self):
        all_clauses = self.audit_engine.get_all_clauses()
        self.results_table.setRowCount(0)

        for clause_number, description, keywords, expected_evidence_text in all_clauses:
            row_position = self.results_table.rowCount()
            self.results_table.insertRow(row_position)

            clause_item = QTableWidgetItem(f"{clause_number} - {description}")
            self.results_table.setItem(row_position, 0, clause_item)

            status_item = QTableWidgetItem("---")
            status_item.setForeground(QColor("gray"))
            self.results_table.setItem(row_position, 1, status_item)

            expected_evidence_item = QTableWidgetItem(expected_evidence_text)
            self.results_table.setItem(row_position, 2, expected_evidence_item)

        self.results_table.resizeRowsToContents()
        self.clause_details_label.setText("برای مشاهده جزئیات بیشتر، روی بند کلیک کنید.")
        self.clause_description_text.clear()
        self.search_results_text.clear()

    def search_in_content(self):
        query = self.search_input.text().strip()
        if not query:
            self.search_results_text.setHtml("<p style='color:red;'>لطفاً کلمه یا عبارت مورد جستجو را وارد کنید.</p>")
            return

        if not self.uploaded_files:
            self.search_results_text.setHtml("<p style='color:red;'>هیچ فایلی برای جستجو بارگذاری نشده است.</p>")
            return

        found_results = {}
        for file_path, content in self.uploaded_files.items():
            if content:
                lower_content = content.lower()
                lower_query = query.lower()

                matches = [(m.start(), m.end()) for m in re.finditer(re.escape(lower_query), lower_content)]
                if matches:
                    found_results[file_path] = []
                    for start_idx, end_idx in matches:
                        context_length = 150
                        start = max(0, start_idx - context_length)
                        end = min(len(content), end_idx + context_length)
                        excerpt = content[start:end]

                        highlighted_excerpt = excerpt.replace(query, f"<span style='background-color: yellow; font-weight: bold;'>{query}</span>", 1)
                        if query.lower() not in highlighted_excerpt.lower():
                            highlighted_excerpt = highlighted_excerpt.replace(query.lower(), f"<span style='background-color: yellow; font-weight: bold;'>{query.lower()}</span>", 1)

                        found_results[file_path].append(f"... {highlighted_excerpt} ...")

        if found_results:
            results_html = "<h3>نتایج جستجو:</h3>"
            for file_path, excerpts in found_results.items():
                results_html += f"<p style='color: #0056b3; font-weight: bold;'>فایل: {os.path.basename(file_path)}</p>"
                for excerpt in excerpts:
                    results_html += f"<p style='margin-right: 20px;'>{excerpt}</p><hr>"
            self.search_results_text.setHtml(results_html)
        else:
            self.search_results_text.setHtml(f"<p style='color: #dc3545;'>کلمه '{query}' در هیچ یک از مستندات یافت نشد.</p>")

    def show_report_options(self):
        if not self.current_audit_results:
            QMessageBox.warning(self, "خطا در گزارش‌گیری", "لطفاً ابتدا ممیزی را انجام دهید تا نتایج برای گزارش‌گیری موجود باشد.")
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("انتخاب نوع گزارش")
        dialog.setFixedSize(300, 150)
        
        layout = QVBoxLayout()
        label = QLabel("لطفاً نوع گزارش مورد نظر را انتخاب کنید:")
        label.setFont(self.main_font)
        layout.addWidget(label, alignment=Qt.AlignCenter)

        png_button = QPushButton("PNG (تصویر جدول نتایج)")
        png_button.setFont(self.main_font)
        png_button.setStyleSheet("background-color: #6c757d; color: white; padding: 8px; border-radius: 4px;")
        png_button.clicked.connect(lambda: self.generate_png_report(dialog))
        layout.addWidget(png_button)

        word_button = QPushButton("WORD (گزارش متنی جامع)")
        word_button.setFont(self.main_font)
        word_button.setStyleSheet("background-color: #007bff; color: white; padding: 8px; border-radius: 4px;")
        word_button.clicked.connect(lambda: self.generate_word_report(dialog))
        layout.addWidget(word_button)

        dialog.setLayout(layout)
        dialog.exec_()

    def generate_png_report(self, parent_dialog):
        parent_dialog.accept() # بستن دیالوگ انتخاب نوع گزارش

        if not self.results_table.rowCount():
            QMessageBox.warning(self, "خطا", "جدول نتایج ممیزی خالی است. نمی‌توان گزارش تصویری تهیه کرد.")
            return

        pixmap = self.results_table.grab()

        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, "ذخیره گزارش PNG", "گزارش_ممیزی_ISO45001.png", "PNG Files (*.png);;All Files (*)", options=options)

        if file_name:
            try:
                pixmap.save(file_name, "PNG")
                QMessageBox.information(self, "موفقیت", f"گزارش تصویری با موفقیت در '{file_name}' ذخیره شد.")
            except Exception as e:
                QMessageBox.critical(self, "خطا", f"خطا در ذخیره گزارش PNG: {str(e)}")

    def generate_word_report(self, parent_dialog):
        parent_dialog.accept() # بستن دیالوگ انتخاب نوع گزارش

        if not self.current_audit_results:
            QMessageBox.warning(self, "خطا در گزارش‌گیری", "هیچ نتیجه ممیزی برای ایجاد گزارش Word وجود ندارد.")
            return

        doc = Document()
        doc.add_heading('گزارش ممیزی سیستم مدیریت ایمنی و بهداشت شغلی ISO 45001:2018', level=1)
        
        
        from datetime import datetime
        report_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"تاریخ گزارش: {report_date}")

        doc.add_heading('خلاصه نتایج ممیزی', level=2)

        total_clauses = len(self.current_audit_results)
        conformant_clauses = 0
        minor_non_conformities = 0
        major_non_conformities = 0

        for clause_num, data in self.current_audit_results.items():
            if not data['non_conformities']:
                conformant_clauses += 1
            else:
                if any(nc['type'].startswith("Major") for nc in data['non_conformities']):
                    major_non_conformities += 1
                elif any(nc['type'].startswith("Minor") for nc in data['non_conformities']):
                    minor_non_conformities += 1
                
        doc.add_paragraph(f"تعداد کل بندهای ممیزی شده: {total_clauses}")
        doc.add_paragraph(f"تعداد بندهای با انطباق کامل: {conformant_clauses}")
        doc.add_paragraph(f"تعداد بندهای با عدم انطباق جزئی: {minor_non_conformities}")
        doc.add_paragraph(f"تعداد بندهای با عدم انطباق اصلی: {major_non_conformities}")

        doc.add_page_break()
        doc.add_heading('جزئیات بند به بند ممیزی', level=2)

        for clause_num, data in self.current_audit_results.items():
            doc.add_heading(f"بند {clause_num}: {data['description']}", level=3)

            status_text = "انطباق کامل"
            if data['non_conformities']:
                if any(nc['type'].startswith("Major") for nc in data['non_conformities']):
                    status_text = "عدم انطباق اصلی"
                elif any(nc['type'].startswith("Minor") for nc in data['non_conformities']):
                    status_text = "عدم انطباق جزئی"
                    
            doc.add_paragraph(f"وضعیت انطباق: {status_text}")

            
            expected_evidence_text = data.get('expected_evidence', "اطلاعات شواهد مورد انتظار موجود نیست.")
            doc.add_paragraph(f"شواهد مورد انتظار: {expected_evidence_text}")

            if data['found_keywords']:
                doc.add_paragraph(f"کلمات کلیدی مرتبط یافت شده: {', '.join(data['found_keywords'])}")
            else:
                doc.add_paragraph("کلمات کلیدی مرتبط یافت نشدند.")

            if data['non_conformities']:
                doc.add_paragraph("عدم انطباق‌های احتمالی:")
                for nc in data['non_conformities']:
                    doc.add_paragraph(f"  - نوع: {nc['type']}")
                    doc.add_paragraph(f"    فایل مرتبط: {os.path.basename(nc['file_path'])}")
                    doc.add_paragraph(f"    بخشی از متن: \"{nc['excerpt']}\"")
            doc.add_paragraph("\n") # یک پاراگراف خالی برای جداسازی بندها

        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, "ذخیره گزارش Word", "گزارش_ممیزی_ISO45001.docx", "Word Documents (*.docx);;All Files (*)", options=options)

        if file_name:
            try:
                doc.save(file_name)
                QMessageBox.information(self, "موفقیت", f"گزارش Word با موفقیت در '{file_name}' ذخیره شد.")
            except Exception as e:
                QMessageBox.critical(self, "خطا", f"خطا در ذخیره گزارش Word: {str(e)}")
    def request_db_password(self):
        default_password = "Hhs@24121365" # رمز عبور پیش فرض

        password, ok = QInputDialog.getText(self,
                                            "ورود به مدیریت پایگاه دانش",
                                            "لطفاً رمز عبور را وارد کنید:",
                                            QLineEdit.Password) # QLineEdit.Password برای پنهان کردن کاراکترها

        if ok: 
            if password == default_password:
                QMessageBox.information(self, "موفقیت", "رمز عبور صحیح است. دسترسی به مدیریت پایگاه دانش.")
                self._open_db_manager() # فراخوانی متد باز کردن دیالوگ مدیریت پایگاه دانش
            else:
                QMessageBox.critical(self, "خطا", "رمز عبور اشتباه است.")
        else: 
            QMessageBox.information(self, "لغو", "ورود لغو شد.")
    def _open_db_manager(self):
        db_dialog = DbManagerDialog(self)
        db_dialog.exec_()


class DbManagerDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("مدیریت پایگاه دانش ISO 45001")
        self.setMinimumSize(900, 600)

        self.db_path = "audit_data.db"

        self.init_ui()
        self.load_data()

    def init_ui(self):
        main_layout = QVBoxLayout()

        
        self.table_widget = QTableWidget()
        
        self.table_widget.setColumnCount(5) # clause_id, clause_number, description, keywords, expected_evidence
        self.table_widget.setHorizontalHeaderLabels([
            "ID بند", "شماره بند", "توضیحات بند", "کلمات کلیدی", "شواهد مورد انتظار"
        ])
        self.table_widget.horizontalHeader().setStretchLastSection(True)
        self.table_widget.setEditTriggers(QTableWidget.NoEditTriggers)
        self.table_widget.itemSelectionChanged.connect(self.display_selected_clause)

        main_layout.addWidget(self.table_widget)

        
        input_layout = QVBoxLayout()
        self.clause_id_label = QLabel("ID بند (فقط برای ویرایش/حذف):")
        self.clause_id_input = QLineEdit()
        self.clause_id_input.setReadOnly(True)
        input_layout.addWidget(self.clause_id_label)
        input_layout.addWidget(self.clause_id_input)

        self.clause_number_label = QLabel("شماره بند:")
        self.clause_number_input = QLineEdit()
        input_layout.addWidget(self.clause_number_label)
        input_layout.addWidget(self.clause_number_input)

        self.description_label = QLabel("توضیحات بند:")
        self.description_input = QTextEdit()
        self.description_input.setFixedHeight(60)
        input_layout.addWidget(self.description_label)
        input_layout.addWidget(self.description_input)

        self.keywords_label = QLabel("کلمات کلیدی (با کاما جدا کنید):")
        self.keywords_input = QTextEdit()
        self.keywords_input.setFixedHeight(60)
        input_layout.addWidget(self.keywords_label)
        input_layout.addWidget(self.keywords_input)

        self.evidence_label = QLabel("شواهد مورد انتظار:")
        self.evidence_input = QTextEdit()
        self.evidence_input.setFixedHeight(60)
        input_layout.addWidget(self.evidence_label)
        input_layout.addWidget(self.evidence_input)

        main_layout.addLayout(input_layout)

        
        button_layout = QHBoxLayout()
        self.add_button = QPushButton("افزودن بند جدید")
        self.add_button.setStyleSheet("background-color: #28a745; color: white;")
        self.add_button.clicked.connect(self.add_clause)
        button_layout.addWidget(self.add_button)

        self.update_button = QPushButton("به‌روزرسانی بند")
        self.update_button.setStyleSheet("background-color: #007bff; color: white;")
        self.update_button.clicked.connect(self.update_clause)
        button_layout.addWidget(self.update_button)

        self.delete_button = QPushButton("حذف بند")
        self.delete_button.setStyleSheet("background-color: #dc3545; color: white;")
        self.delete_button.clicked.connect(self.delete_clause)
        button_layout.addWidget(self.delete_button)

        self.clear_button = QPushButton("پاک کردن فیلدها")
        self.clear_button.clicked.connect(self.clear_fields)
        button_layout.addWidget(self.clear_button)

        main_layout.addLayout(button_layout)
        self.setLayout(main_layout)

    def load_data(self):
        self.table_widget.setRowCount(0)
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT clause_id, clause_number, description, keywords, expected_evidence FROM iso_standards")
        rows = cursor.fetchall()
        conn.close()

        self.table_widget.setRowCount(len(rows))
        for row_idx, row_data in enumerate(rows):
            for col_idx, col_data in enumerate(row_data):
                item = QTableWidgetItem(str(col_data))
                self.table_widget.setItem(row_idx, col_idx, item)
        
        self.table_widget.setColumnWidth(0, 50)  # ID
        self.table_widget.setColumnWidth(1, 100) # شماره بند
        self.table_widget.setColumnWidth(2, 250) # توضیحات
        self.table_widget.setColumnWidth(3, 150) # کلمات کلیدی
        self.table_widget.setColumnWidth(4, 250) # شواهد مورد انتظار
        self.table_widget.resizeRowsToContents()


    def display_selected_clause(self):
        selected_items = self.table_widget.selectedItems()
        if selected_items:
            row = selected_items[0].row()
            self.clause_id_input.setText(self.table_widget.item(row, 0).text())
            self.clause_number_input.setText(self.table_widget.item(row, 1).text())
            self.description_input.setText(self.table_widget.item(row, 2).text())
            self.keywords_input.setText(self.table_widget.item(row, 3).text())
            
            if self.table_widget.columnCount() > 4: 
                self.evidence_input.setText(self.table_widget.item(row, 4).text())
            else:
                self.evidence_input.clear() 

    def add_clause(self):
        clause_number = self.clause_number_input.text().strip()
        description = self.description_input.toPlainText().strip()
        keywords = self.keywords_input.toPlainText().strip()
        expected_evidence = self.evidence_input.toPlainText().strip()

        if not clause_number or not description or not keywords:
            QMessageBox.warning(self, "ورودی ناقص", "لطفاً شماره بند، توضیحات و کلمات کلیدی را وارد کنید.")
            return

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        try:
            cursor.execute(
                "INSERT INTO iso_standards (clause_number, description, keywords, expected_evidence) VALUES (?, ?, ?, ?)",
                (clause_number, description, keywords, expected_evidence)
            )
            conn.commit()
            QMessageBox.information(self, "موفقیت", "بند جدید با موفقیت اضافه شد.")
            self.load_data()
            self.clear_fields()
        except sqlite3.IntegrityError:
            QMessageBox.warning(self, "خطا", "شماره بند تکراری است. لطفاً یک شماره بند منحصر به فرد وارد کنید.")
        except Exception as e:
            QMessageBox.critical(self, "خطا", f"خطا در افزودن بند: {e}")
        finally:
            conn.close()

    def update_clause(self):
        clause_id = self.clause_id_input.text().strip()
        clause_number = self.clause_number_input.text().strip()
        description = self.description_input.toPlainText().strip()
        keywords = self.keywords_input.toPlainText().strip()
        expected_evidence = self.evidence_input.toPlainText().strip()

        if not clause_id:
            QMessageBox.warning(self, "انتخاب بند", "لطفاً بندی را از جدول برای به‌روزرسانی انتخاب کنید.")
            return
        if not clause_number or not description or not keywords:
            QMessageBox.warning(self, "ورودی ناقص", "لطفاً شماره بند، توضیحات و کلمات کلیدی را وارد کنید.")
            return

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        try:
            cursor.execute(
                "UPDATE iso_standards SET clause_number = ?, description = ?, keywords = ?, expected_evidence = ? WHERE clause_id = ?",
                (clause_number, description, keywords, expected_evidence, clause_id)
            )
            conn.commit()
            if cursor.rowcount > 0:
                QMessageBox.information(self, "موفقیت", "بند با موفقیت به‌روزرسانی شد.")
                self.load_data()
                self.clear_fields()
            else:
                QMessageBox.warning(self, "خطا", "بند یافت نشد یا تغییر اعمال نشد.")
        except sqlite3.IntegrityError:
            QMessageBox.warning(self, "خطا", "شماره بند تکراری است. لطفاً یک شماره بند منحصر به فرد وارد کنید.")
        except Exception as e:
            QMessageBox.critical(self, "خطا", f"خطا در به‌روزرسانی بند: {e}")
        finally:
            conn.close()

    def delete_clause(self):
        clause_id = self.clause_id_input.text().strip()

        if not clause_id:
            QMessageBox.warning(self, "انتخاب بند", "لطفاً بندی را از جدول برای حذف انتخاب کنید.")
            return

        reply = QMessageBox.question(self, 'تأیید حذف', f"آیا مطمئنید که می‌خواهید بند با ID {clause_id} را حذف کنید؟",
                                     QMessageBox.Yes | QMessageBox.No, QMessageBox.No)

        if reply == QMessageBox.Yes:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            try:
                cursor.execute("DELETE FROM iso_standards WHERE clause_id = ?", (clause_id,))
                conn.commit()
                if cursor.rowcount > 0:
                    QMessageBox.information(self, "موفقیت", "بند با موفقیت حذف شد.")
                    self.load_data()
                    self.clear_fields()
                else:
                    QMessageBox.warning(self, "خطا", "بند یافت نشد یا حذف نشد.")
            except Exception as e:
                QMessageBox.critical(self, "خطا", f"خطا در حذف بند: {e}")
            finally:
                conn.close()

    def clear_fields(self):
        
        self.clause_id_input.clear()
        self.clause_number_input.clear()
        self.description_input.clear()
        self.keywords_input.clear()
        self.evidence_input.clear()
        self.table_widget.clearSelection() 


def create_db_if_not_exists():
    conn = sqlite3.connect("audit_data.db")
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS iso_standards (
            clause_id INTEGER PRIMARY KEY AUTOINCREMENT,
            clause_number TEXT UNIQUE NOT NULL,
            description TEXT,
            keywords TEXT,
            expected_evidence TEXT
        )
    """)
    conn.commit()
    conn.close()


if __name__ == "__main__":
    # در ابتدا، مطمئن شوید که پایگاه داده و جدول وجود دارند و به روز هستند
    create_db_if_not_exists()
    
    app = QApplication(sys.argv)
    window = MainWindow()
    window.showMaximized() # یا window.show()
    sys.exit(app.exec_())